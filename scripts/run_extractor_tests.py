import sys
from pathlib import Path
import sys as _sys
# ensure project src is importable
_sys.path.insert(0, str(Path(__file__).resolve().parents[1] / 'src'))
from extractor import Extractor

SAMPLES_DIR = Path("tests/samples")
SAMPLES_DIR.mkdir(parents=True, exist_ok=True)

AGREE1 = """
1. Definitions
This Agreement contains definitions used throughout.

2. Term
The term begins on the Effective Date.

3. Termination
Either party may terminate.
"""
(SAMPLES_DIR / "agreement1.txt").write_text(AGREE1, encoding='utf-8')

AGREE2 = """
Definitions
These are the definitions.

Warranty
Supplier warrants the goods.

Limitation of Liability
No liability beyond fees.
"""
(SAMPLES_DIR / "agreement2.txt").write_text(AGREE2, encoding='utf-8')

# docx sample if possible
try:
    from docx import Document
    doc = Document()
    doc.add_paragraph("Confidentiality")
    doc.add_paragraph("The parties agree to keep information secret.")
    doc.add_paragraph("Governing Law")
    doc.add_paragraph("This Agreement is governed by State law.")
    doc.save(SAMPLES_DIR / "agreement3.docx")
    have_docx = True
except Exception:
    have_docx = False

ex = Extractor()

print('Running test 1: numbered txt')
out1 = ex.extract_from_file(str(SAMPLES_DIR / 'agreement1.txt'))
print('EXTRACTED CLAUSES 1:', out1)
if not any('definitions' in c['clause_type'] for c in out1):
    print('FAILED: definitions not found in agreement1')
    sys.exit(2)
if not any('termination' in c['clause_type'] for c in out1):
    print('FAILED: termination not found in agreement1')
    sys.exit(2)
print('OK')

print('Running test 2: title-case txt')
out2 = ex.extract_from_file(str(SAMPLES_DIR / 'agreement2.txt'))
if not any('warranty' in c['clause_type'] for c in out2):
    print('FAILED: warranty not found in agreement2')
    sys.exit(2)
if not any('limitation' in c['clause_type'] for c in out2):
    print('FAILED: limitation not found in agreement2')
    sys.exit(2)
print('OK')

if have_docx:
    print('Running test 3: docx')
    out3 = ex.extract_from_file(str(SAMPLES_DIR / 'agreement3.docx'))
    if not any('confidentiality' in c['clause_type'] for c in out3):
        print('FAILED: confidentiality not found in agreement3')
        sys.exit(2)
    if not any('governing' in c['clause_type'] for c in out3):
        print('FAILED: governing law not found in agreement3')
        sys.exit(2)
    print('OK')
else:
    print('Skipping docx test; python-docx not available')

print('All checks passed')
sys.exit(0)
